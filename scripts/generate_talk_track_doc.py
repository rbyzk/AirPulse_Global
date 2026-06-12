from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "AirPulse_Global_Sunum_Konusma_Metni.docx"

BLUE = RGBColor(47, 99, 216)
DARK = RGBColor(22, 28, 45)
MUTED = RGBColor(92, 99, 112)


SLIDES: list[tuple[str, str]] = [
    (
        "Slayt 1 - Açılış",
        "Merhaba, ben Beyza. Bugün size AirPulse Global projesini anlatacağım. Bu projeyi yapma nedenim, hava kalitesi verilerinin çoğu zaman yalnızca teknik tablolar ve sayılar halinde karşımıza çıkmasıydı. İnsanlar çoğu zaman havanın kirli olduğunu görüyor ama bunun kendileri için ne anlama geldiğini, günlük yaşamlarında neyi değiştirmeleri gerektiğini tam olarak anlayamıyor. Ben de bu noktada, çevresel veriyi daha anlaşılır, daha kişisel ve daha eyleme dönük hale getiren bir proje geliştirmek istedim. AirPulse Global'in temel amacı, hava kalitesi verisini sadece göstermek değil; onu yorumlamak, kullanıcıya uygun hale getirmek ve karar vermeyi kolaylaştırmak.",
    ),
    (
        "Slayt 2 - Vizyon",
        "Bu slaytta projenin üç temel vizyonunu görüyoruz. Birincisi demokratikleştirme; yani teknik ve karmaşık çevre verilerini herkesin birkaç saniye içinde anlayabileceği bir dile çevirmek. İkincisi kişiselleştirme; kullanıcıya genel bir rapor vermek yerine bulunduğu konum, rüzgar bilgisi ve çevresel bağlama göre daha anlamlı öneriler sunmak. Üçüncüsü ise sürdürülebilirlik; yani sadece problemi göstermek değil, bireysel davranışları daha çevreci hale getirecek bir aksiyon katmanı oluşturmak. Kısacası bu proje, veriden farkındalığa, farkındalıktan da davranış değişikliğine giden bir köprü kurmayı amaçlıyor.",
    ),
    (
        "Slayt 3 - Problem ve Fırsat",
        "AirPulse’un çıkış noktası üç temel soruna dayanıyor. İlk olarak veriler parçalı halde bulunuyor; farklı istasyonlar, farklı veri akışları ve farklı hava durumu servisleri var. İkinci olarak bu veriler teknik olduğu için son kullanıcı tarafından kolay yorumlanamıyor. Üçüncü olarak ise kullanıcı çoğu zaman sayıyı görüyor ama ne yapması gerektiğini bilmiyor. İşte bu noktada AirPulse fırsat alanı oluşturuyor. İstanbul odaklı başlayan bu yapı, zamanla daha geniş kapsama sahip, global ölçekte çalışabilen bir çevresel karar destek platformuna dönüşüyor.",
    ),
    (
        "Slayt 4 - Kullanıcı Yolculuğu",
        "Burada kullanıcının ürün içindeki temel yolculuğunu görüyoruz. Süreç şehir ya da istasyon seçimiyle başlıyor. Sonrasında kullanıcı hava kalitesini görüyor, istasyonları karşılaştırıyor, kısa vadeli tahmini inceliyor ve en sonunda aksiyon ile raporlama katmanına geçiyor. Bu akışın önemli tarafı şu: AirPulse yalnızca anlık AQI (Air Quality Index) gösteren bir uygulama değil, veriyi içgörüye, içgörüyü de davranışa dönüştüren bütünleşik bir akış sunuyor.",
    ),
    (
        "Slayt 5 - Sistem Mimarisi ve Veri Akışı",
        "Bu slaytta mimariyi ve veri akışını özetliyoruz. Girdi tarafında üç ana kaynak var: WAQI (vey-ki) canlı hava kalitesi verisi, Open-Meteo (open metio) destekleyici veri akışı ve Tomorrow.io (tumoro aiyo) üzerinden gelen rüzgar bağlamı. Bu veriler uygulama içinde normalize ediliyor, birleştiriliyor ve istasyon bazlı olarak anlamlandırılıyor. Sonraki aşamada tahmin, analitik çıkarımlar ve sağlık rehberliği üretiliyor. En sonunda bu bilgiler dashboard, raporlar ve dışa aktarım çıktıları olarak kullanıcıya sunuluyor. Bu yapı, veri toplama ile kullanıcı deneyimi arasındaki zincirin kopmadan ilerlemesini sağlıyor.",
    ),
    (
        "Slayt 6 - Teknoloji Yığını",
        "Projenin teknoloji tarafında hızlı ürün geliştirme ve güçlü veri işleme dengesini kurmaya çalıştım. Arayüz ve uygulama akışı için Streamlit kullanıldı. Veri işleme ve dönüştürme tarafında Pandas ve NumPy yer alıyor. Görselleştirme için Plotly ve harita tarafında Folium kullanıldı. Raporlama çıktılarında ise ReportLab ve Pillow tercih edildi. Bu seçimlerin ortak amacı, canlı veriyi hızlı şekilde görselleştirmek, kullanıcıya anlaşılır sunmak ve çıktı alınabilir hale getirmek oldu.",
    ),
    (
        "Slayt 7 - Veri Kaynakları",
        "Burada kullandığımız temel veri kaynaklarını görüyoruz. WAQI (vey-ki), ana hava kalitesi akışı ve istasyon seviyesinde forecast (forkast) kaynağı olarak görev yapıyor. Tomorrow.io (tumoro aiyo), özellikle rüzgar yönü, hızı ve gust bilgisini vererek dağılım ve maruziyet yorumlarını güçlendiriyor. Open-Meteo (open metio) ise anahtarsız, global ve yardımcı bir veri kaynağı olarak sistemi daha dayanıklı hale getiriyor. Böylece tek bir kaynağa bağlı kalmadan daha esnek bir mimari kurulmuş oluyor.",
    ),
    (
        "Slayt 8 - İzlenen Kirleticiler",
        "AirPulse farklı kirletici türlerini birlikte ele alıyor. Burada geçen parametreleri kısaca açıklamak istiyorum. PM2.5 (Particulate Matter 2.5 - pi-em iki nokta beş), havada bulunan çok küçük partikül maddeleri ifade ediyor ve özellikle akciğerlerin derinlerine kadar ulaşabildiği için sağlık açısından oldukça kritik. PM10 (Particulate Matter 10), yine partikül madde ama biraz daha büyük parçacıkları temsil ediyor. NO2 (Nitrogen Dioxide - en-o-iki), çoğunlukla trafik kaynaklı azot dioksiti ifade ediyor. SO2 (Sulfur Dioxide - es-o-iki), özellikle sanayi ve yakıt kullanımına bağlı kükürt dioksiti gösteriyor. O3 (Ozone - o-üç), yer seviyesindeki ozonu ifade ediyor. CO (Carbon Monoxide - si-o) ise karbon monoksiti temsil ediyor. Buradaki amaç sadece bu kısaltmaları göstermek değil; kullanıcının hangi kirleticinin ne anlama geldiğini, hangi kaynaklardan çıkabildiğini ve neden önemli olduğunu anlayabilmesini sağlamak.",
    ),
    (
        "Slayt 9 - Tahminleme ve Doğrulama",
        "Tahmin tarafında öncelik resmi WAQI günlük forecast verisinde. Eğer bu veri o şehir veya istasyon için mevcut değilse sistem daha muhafazakar bir fallback, yani geri dönüş tahmin yaklaşımına geçiyor. Bu slaytta özellikle model tarafının yalnızca teorik kalmadığını göstermek istedim. Örneğin offline PM2.5 champion modelinde R kare değeri 0.77, RMSE (ar-em-es-i) değeri ise 16.27 seviyesinde. Ayrıca cross-station benchmark tarafında 8 binin üzerinde satırla değerlendirme yapılmış durumda. Bunun anlamı şu: sistem yalnızca tahmin üretiyor demiyoruz, aynı zamanda bu tahminleri izliyor, kıyaslıyor ve ne kadar güvenilebilir olduğunu da göstermeye çalışıyoruz.",
    ),
    (
        "Slayt 10 - Ürün Ekranları",
        "Burada ürünün iki ana etkileşim alanını görüyoruz. Solda dashboard yapısı, yani kullanıcının ilk karşılaştığı karar ekranı var. Bu bölümde AQI (Air Quality Index), genel sağlık durumu, rüzgar bilgisi ve hızlı yorumlar bir arada sunuluyor. Sağ tarafta ise istasyon haritası yer alıyor. Bu alan, şehir ortalamasının ötesine geçip mahalle veya istasyon düzeyinde farkları görmeye yardımcı oluyor. Kullanıcı açısından bu çok önemli çünkü aynı şehir içinde maruziyet seviyesi ciddi biçimde değişebiliyor.",
    ),
    (
        "Slayt 11 - Aksiyon Katmanı",
        "Bu proje yalnızca izleme amacı taşımıyor; kullanıcıyı davranış değişikliğine yönlendirmeye çalışıyor. Bu nedenle karbon ayak izi hesaplama, ulaşım tercihi karşılaştırması ve günlük checklist gibi bileşenler eklendi. Kullanıcı, çevresel veriyi soyut bir bilgi olarak değil, kendi yaşamına temas eden bir karar alanı olarak görsün istedim. Örneğin bugün araç yerine yürümek veya bisiklet tercih etmek hem maruziyet hem karbon etkisi açısından değerlendirilebiliyor.",
    ),
    (
        "Slayt 12 - Analitik Derinlik",
        "Bu bölümde sistemin daha analitik okuma tarafını görüyoruz. Rüzgar ile kirletici yoğunluğu arasındaki ilişki, kaynak tipi tahminleri, forecast doğrulama ve anomali tespiti burada öne çıkıyor. Bu katman, uygulamanın sadece görsel bir panel olmadığını; veri üzerinde yorum üreten bir analiz aracı olduğunu göstermesi açısından önemli. Özellikle anomali tespiti ve validation tarafı, projenin ileride daha güçlü karar destek mekanizmalarına dönüşebileceğini gösteriyor.",
    ),
    (
        "Slayt 13 - Raporlama",
        "AirPulse’un güçlü yönlerinden biri de çıktılarının paylaşılabilir olması. A4 PDF rapor üretimi, sosyal medya için görsel kartlar ve CSV (si-es-vi) dışa aktarımı sayesinde hem teknik hem iletişimsel kullanım destekleniyor. Bu da projeyi yalnızca bireysel kullanıcı değil; okul, kurum, belediye veya farkındalık kampanyaları gibi senaryolar için de daha anlamlı hale getiriyor.",
    ),
    (
        "Slayt 14 - Sonuçlar ve Kanıtlar",
        "Bu slayt projeyi savunurken en kritik bölümlerden biri. Çünkü burada yalnızca özellik anlatmıyoruz, aynı zamanda bu özelliklerin arkasındaki kanıtları gösteriyoruz. Şu anda 25’ten fazla global şehir ön ayarı, 3 canlı veri sağlayıcısı ve 3 ila 10 gün arasında değişen tahmin ufku bulunuyor. Ayrıca 48 adet önbelleğe alınmış şehir veya istasyon geçmişi sayesinde tekrar edilebilir demolar yapılabiliyor. PM2.5 için offline champion model sonucu, leave-station-out benchmark sonucu ve validation store kayıtları; sistemin hem ürün hem de değerlendirme tarafında olgunlaştığını gösteriyor. Yani bu proje sadece iyi görünen bir arayüz değil, arkasında ölçülmüş ve takip edilmiş bir teknik altyapı taşıyor.",
    ),
    (
        "Slayt 15 - Demo Odağı, Yol Haritası ve Limitler",
        "Burada iki şeyi birlikte anlatmak önemli. Birincisi, demo sırasında en güçlü akış şehir arama, forecast inceleme, take action ve rapor export zinciri. Yani en etkileyici kullanıcı hikayesi bu akışta ortaya çıkıyor. İkincisi ise limitlerimizi dürüst biçimde ifade etmek. Sistem bazı özelliklerde API bağımlılığı taşıyor, rüzgar tarafı Tomorrow.io yapılandırıldığında daha güçlü çalışıyor ve şehir bazlı kapsam farklılıkları olabiliyor. Yol haritasında ise bildirimler, daha güçlü validasyon setleri ve kurumsal raporlama görünümü gibi genişleme alanları bulunuyor.",
    ),
    (
        "Slayt 16 - Demo Geçişi ve Kapanış",
        "Bu noktada sunum tarafında genel yapıyı özetlemiş oldum. Şimdi isterseniz Streamlit (strimlit) sayfası üzerinden projenin gerçek arayüzde nasıl göründüğüne bakalım. Böylece anlattığım veri akışının, dashboard yapısının, tahmin katmanının ve kullanıcı aksiyon alanlarının uygulama içinde nasıl çalıştığını daha somut şekilde görebiliriz. Ardından son olarak teşekkür kısmına geçebilirim. Dinlediğiniz için teşekkür ederim, sorularınızı memnuniyetle cevaplayabilirim.",
    ),
]


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = BLUE


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = BLUE


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(11)
    run.font.color.rgb = DARK


def build_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.25)
    section.right_margin = Cm(2.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)

    add_title(doc, "AirPulse Global Sunum Konuşma Metni")
    add_subtitle(doc, "Slayt sırasına göre akıcı, anlaşılır ve sunum sırasında doğrudan kullanılabilecek Türkçe konuşma metni")

    for heading, body in SLIDES:
        add_heading(doc, heading)
        add_body(doc, body)

    doc.save(OUTPUT)
    print(f"Saved talk track to: {OUTPUT}")


if __name__ == "__main__":
    build_doc()
