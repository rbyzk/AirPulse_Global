from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "AirPulse_Global_Streamlit_Demo_Metni.docx"

BLUE = RGBColor(47, 99, 216)
DARK = RGBColor(22, 28, 45)
MUTED = RGBColor(92, 99, 112)


SECTIONS: list[tuple[str, str]] = [
    (
        "Açılış",
        "Şimdi sunum tarafını tamamladıktan sonra projenin Streamlit (strimlit) arayüzünde nasıl göründüğünü göstermek istiyorum. Burada amacım yalnızca ekranları göstermek değil; kullanıcı uygulamaya girdiğinde hangi akıştan geçiyor, hangi bilgileri görüyor ve bu bilgiler nasıl aksiyona dönüşüyor, bunu somut olarak anlatmak.",
    ),
    (
        "Ana Sayfa ve Genel Bakış",
        "Öncelikle ana ekranda kullanıcıyı genel bir hava kalitesi özeti karşılıyor. Bu bölümde seçilen şehre ait güncel AQI (Air Quality Index) değeri, temel kirleticiler ve genel sağlık yorumu yer alıyor. Ben burada özellikle şunu vurgulayabilirim: Kullanıcı teknik veriyi ham haliyle görmek zorunda kalmıyor; uygulama bu veriyi daha anlaşılır bir yapıya dönüştürüyor.",
    ),
    (
        "Şehir Arama Akışı",
        "Şimdi şehir arama alanına geliyorum. Burada kullanıcı farklı şehirler arasında geçiş yapabiliyor. Bu önemli çünkü proje tek bir bölgeye bağlı değil; farklı şehirlerdeki hava kalitesi bağlamını karşılaştırmalı şekilde incelemeyi mümkün kılıyor. Ben bu kısımda örnek olarak bir şehir seçip, verinin anlık olarak nasıl güncellendiğini gösterebilirim.",
    ),
    (
        "AQI ve Kirletici Kartları",
        "Seçilen şehirle birlikte AQI (Air Quality Index) değeri ve kirletici kartları güncelleniyor. Burada PM2.5, PM10, NO2, SO2, O3 ve CO gibi kirleticilerin seviyeleri görülebiliyor. Bu bölümü anlatırken özellikle PM2.5'in çok küçük partikül maddeleri temsil ettiğini ve sağlık açısından kritik olduğunu, NO2'nin ise daha çok trafik kaynaklı bir gösterge olduğunu vurgulayabilirim. Böylece kullanıcı için sadece sayı değil, anlam da üretmiş oluyoruz.",
    ),
    (
        "Rüzgar Bağlamı",
        "Bu bölümde rüzgar verisinin neden önemli olduğunu göstermek istiyorum. Çünkü hava kalitesi sadece kirletici miktarıyla değil, rüzgarın yönü ve hızıyla birlikte daha doğru yorumlanabiliyor. Eğer Tomorrow.io (tumoro aiyo) verisi aktifse burada daha güçlü bir bağlam sağlanıyor. Ben bunu anlatırken, aynı kirletici değerinin farklı rüzgar koşullarında farklı etkiler yaratabileceğini söyleyebilirim.",
    ),
    (
        "İstasyon Haritası",
        "Şimdi harita tarafına geçiyorum. Bu ekranın önemli tarafı, şehir ortalamasının ötesine geçerek istasyon bazlı farkları göstermesi. Yani aynı şehir içinde farklı bölgelerde hava kalitesi değişebiliyor. Kullanıcı burada istasyonları tek tek inceleyebiliyor, yoğunluk farklarını görebiliyor ve hangi bölgenin daha riskli olduğunu daha net anlayabiliyor.",
    ),
    (
        "Forecast Sayfası",
        "Burada tahmin ekranını görüyoruz. Sistem öncelikle resmi WAQI (vey-ki) forecast verisini kullanıyor. Eğer bu veri eksikse daha muhafazakar bir fallback tahmin yaklaşımına geçiyor. Bu ekranı anlatırken özellikle projenin sadece bugünü göstermediğini, kısa vadede ne olabileceğine dair de yorum üretmeye çalıştığını söylemek önemli. Böylece kullanıcı sadece mevcut durumu değil, bir sonraki adımı da düşünebiliyor.",
    ),
    (
        "Analytics Bölümü",
        "Analytics kısmı projenin daha teknik ve yorumlayıcı yüzünü gösteriyor. Burada zaman içindeki desenler, rüzgar ile kirletici ilişkisi, forecast doğrulama ve anomali tespiti gibi alanlar bulunuyor. Bu bölümü anlatırken, projenin yalnızca bir kullanıcı arayüzü değil, aynı zamanda veri üzerinde analiz yapan bir yapı olduğunu vurgulayabilirim.",
    ),
    (
        "Take Action Bölümü",
        "Bence projenin en ayırt edici taraflarından biri bu bölüm. Çünkü burada kullanıcıya sadece hava kalitesinin kötü olduğunu söylemiyoruz; bunun karşılığında ne yapabileceğini de göstermeye çalışıyoruz. Karbon ayak izi hesabı, ulaşım tercihi karşılaştırması ve günlük checklist gibi bileşenler sayesinde kullanıcı veri ile kendi davranışı arasında bağlantı kurabiliyor.",
    ),
    (
        "Reports ve Dışa Aktarım",
        "Burada raporlama bölümünü gösteriyorum. Kullanıcı bu sistemden PDF rapor, sosyal medya kartı ya da CSV (si-es-vi) çıktısı alabiliyor. Bu özellik önemli çünkü uygulama içindeki veriler yalnızca ekranda kalmıyor; paylaşılabilir, raporlanabilir ve daha profesyonel biçimde sunulabilir hale geliyor.",
    ),
    (
        "Kapanış",
        "Streamlit arayüzünde genel akışa baktığımızda şunu görüyoruz: AirPulse Global veri topluyor, bu veriyi yorumluyor, kullanıcıya uygun hale getiriyor ve en sonunda kullanıcıyı aksiyona yönlendiriyor. Yani bu proje yalnızca hava kalitesini izleyen bir sistem değil; çevresel veriyi karar destek aracına dönüştüren bütünleşik bir platform. Buradan sonra tekrar sunuma dönüp kısa bir kapanış yapabilirim.",
    ),
]


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = BLUE


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = BLUE


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(11)
    run.font.color.rgb = DARK


def build_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.25)
    section.right_margin = Cm(2.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)

    add_title(doc, "AirPulse Global Streamlit Demo Metni")
    add_subtitle(doc, "Canlı uygulama gösterimi sırasında kullanılabilecek Türkçe anlatım metni")

    for heading, body in SECTIONS:
        add_heading(doc, heading)
        add_body(doc, body)

    doc.save(OUTPUT)
    print(f"Saved Streamlit demo script to: {OUTPUT}")


if __name__ == "__main__":
    build_doc()
